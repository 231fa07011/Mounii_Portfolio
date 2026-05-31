import os
import sys

def main():
    try:
        import docx
    except ImportError:
        print("python-docx not found, trying to install it...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    docx_path = "Varikuntla Nagamounika.docx"
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found.")
        return

    print(f"Extracting text from {docx_path}...")
    doc = docx.Document(docx_path)
    
    text = []
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            # Remove duplicate adjacent cells (often caused by merged cells)
            clean_row = []
            for cell in row_text:
                if not clean_row or clean_row[-1] != cell:
                    clean_row.append(cell)
            if any(clean_row):
                text.append(" | ".join(clean_row))

    full_text = "\n".join(text)
    
    with open("resume_docx_text.txt", "w", encoding="utf-8") as f:
        f.write(full_text)
    
    print("Successfully extracted text to resume_docx_text.txt!")

if __name__ == "__main__":
    main()
